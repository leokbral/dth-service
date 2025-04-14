from docx import Document
from docx.oxml.ns import qn
from google.colab import files
import os
import re

# Upload do DOCX
uploaded = files.upload()
docx_path = list(uploaded.keys())[0]

# Abrir documento
doc = Document(docx_path)

# Cria pasta de imagens
os.makedirs("images", exist_ok=True)
image_index = 1
html_output = ""

heading_map = {
    "Heading 1": "h1",
    "Heading 2": "h2",
    "Heading 3": "h3",
    "Heading 4": "h4",
}

bullet_regex = r"^[•\-–‣·◦●◉○\*]\s+"
numbered_regex = r"^\d+[\.\)]\s+"

in_list = False
list_type = None
in_references = False
ref_index = 1
references_html = ""

# Função para salvar imagens embutidas nos parágrafos
def salvar_imagens(run):
    global image_index, html_output
    drawing_elements = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    for blip in drawing_elements:
        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not embed:
            continue
        image_part = run.part.related_parts[embed]
        image_data = image_part.blob
        image_ext = image_part.content_type.split("/")[-1]
        image_name = f"images/image-{image_index}.{image_ext}"
        with open(image_name, "wb") as f:
            f.write(image_data)
        html_output += f'<img src="{image_name}" alt="Imagem {image_index}" style="max-width:100%;"/>\n'
        image_index += 1

# Regex para citações no corpo do texto
def substituir_citacoes(texto):
    return re.sub(r"\[(\d+)\]", r'<a href="#ref-\1">[\1]</a>', texto)

# Processar parágrafos e imagens
for para in doc.paragraphs:
    raw_text = para.text.strip()
    if not raw_text and not any('graphic' in run._element.xml for run in para.runs):
        continue

    style = para.style.name
    text = substituir_citacoes(raw_text)
    is_list_style = style == "List Paragraph"
    is_bullet = re.match(bullet_regex, raw_text)
    is_numbered = re.match(numbered_regex, raw_text)

    # Verifica e salva imagens
    for run in para.runs:
        salvar_imagens(run)

    # Seção de referências
    if style in heading_map:
        if in_references:
            html_output += "<ol>\n" + references_html + "</ol>\n"
            in_references = False
            references_html = ""
            ref_index = 1

        tag = heading_map[style]
        html_output += f"<{tag}>{text}</{tag}>\n"

        if text.lower() in ["referências", "references"]:
            in_references = True
        continue

    if in_references:
        references_html += f'<li id="ref-{ref_index}">{text}</li>\n'
        ref_index += 1
        continue

    # Listas
    if is_list_style or is_bullet or is_numbered:
        new_list_type = "ul" if is_bullet else "ol" if is_numbered else "ul"
        item_text = re.sub(bullet_regex if is_bullet else numbered_regex, '', text)

        if not in_list or new_list_type != list_type:
            if in_list:
                html_output += f"</{list_type}>\n"
            html_output += f"<{new_list_type}>\n"
            in_list = True
            list_type = new_list_type

        html_output += f"<li>{item_text}</li>\n"
    else:
        if in_list:
            html_output += f"</{list_type}>\n"
            in_list = False
            list_type = None

        html_output += f"<p>{text}</p>\n"

# Fecha seção de referências no final
if in_references:
    html_output += "<ol>\n" + references_html + "</ol>\n"

if in_list:
    html_output += f"</{list_type}>\n"

# Processar tabelas
for table in doc.tables:
    html_output += "<table border='1' style='border-collapse:collapse;'>\n"
    for row in table.rows:
        html_output += "<tr>\n"
        for cell in row.cells:
            cell_text = substituir_citacoes(cell.text.strip()).replace("\n", "<br>")
            html_output += f"<td>{cell_text}</td>\n"
        html_output += "</tr>\n"
    html_output += "</table>\n<br/>\n"

# Salvar HTML final
filename = "docx_completo_formatado.html"
with open(filename, "w", encoding="utf-8") as f:
    f.write(html_output)

# Mostrar prévia
print("Prévia do HTML gerado:\n")
print(html_output[:2000])
files.download(filename)
